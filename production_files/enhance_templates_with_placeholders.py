#!/usr/bin/env python3
"""
Enhance Templates with Placeholders - Replace static text with docxtpl placeholders
"""

import os
import re
from docx import Document

def enhance_template_with_placeholders(template_path, output_path):
    """Add docxtpl placeholders to Word template"""
    
    try:
        print(f"📄 Processing template: {os.path.basename(template_path)}")
        
        # Load the document
        doc = Document(template_path)
        
        # Define replacement mappings (static text → placeholder)
        replacements = {
            # Email placeholders
            'info@Domain': '{{primary_email}}',
            'xxx': '{{primary_email_password}}',
            
            # URL placeholders  
            'https://webmail.kunze-marketing.de': '{{webmail_url}}',
            'webmail.kunze-marketing.de': '{{webmail_url}}',
            'https://webmail.kunze-medien.de': '{{webmail_url}}',
            
            # Server placeholders
            'sslmail.kunze-marketing.de': '{{smtp_server}}',
            'sslmail.kunze-medien.de': '{{smtp_server}}',
            
            # Customer placeholders
            'Kundenname': '{{customer_name}}',
            
            # Generic placeholders for common patterns
            'IMAP = 993': 'IMAP = {{imap_port}}',
            'POP3 = 995': 'POP3 = {{pop_port}}', 
            'SMTP = 465': 'SMTP = {{smtp_port}}',
        }
        
        changes_made = 0
        
        # Process all paragraphs
        for para in doc.paragraphs:
            original_text = para.text
            modified_text = original_text
            
            # Apply replacements
            for old_text, new_text in replacements.items():
                if old_text in modified_text:
                    modified_text = modified_text.replace(old_text, new_text)
                    changes_made += 1
                    print(f"   ✅ Replaced: '{old_text}' → '{new_text}'")
            
            # Update paragraph text if changed
            if modified_text != original_text:
                para.text = modified_text
        
        # Process tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        original_text = para.text
                        modified_text = original_text
                        
                        for old_text, new_text in replacements.items():
                            if old_text in modified_text:
                                modified_text = modified_text.replace(old_text, new_text)
                                changes_made += 1
                                print(f"   ✅ Table: '{old_text}' → '{new_text}'")
                        
                        if modified_text != original_text:
                            para.text = modified_text
        
        # Save enhanced template
        doc.save(output_path)
        
        print(f"✅ Enhanced template saved: {output_path}")
        print(f"📊 Total replacements made: {changes_made}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error enhancing template: {e}")
        return False

def add_additional_placeholders(template_path, output_path, template_type):
    """Add additional placeholders specific to template type"""
    
    try:
        doc = Document(template_path)
        
        # Additional placeholders based on vault data
        additional_replacements = {}
        
        if template_type == 'extern':
            # External customer template - focus on email
            additional_replacements.update({
                'E-Mail-Adresse:': 'E-Mail-Adresse: {{primary_email}}',
                'Passwort:': 'Passwort: {{primary_email_password}}',
                'Webmail-Login:': 'Webmail-Login: {{primary_email}}',
                'Website-Login:': 'Website-Login: {{website_login}}',
            })
        else:
            # Internal customer template - more comprehensive
            additional_replacements.update({
                'Website-Login:': 'Website-Login: {{website_login}}',
                'Website-Passwort:': 'Website-Passwort: {{website_password}}',
                'Statistik-Login:': 'Statistik-Login: {{statistics_login}}',
                'Statistik-Passwort:': 'Statistik-Passwort: {{statistics_password}}',
                'E-Mail 1:': 'E-Mail 1: {{primary_email}}',
                'E-Mail 2:': 'E-Mail 2: {{secondary_email}}',
            })
        
        changes_made = 0
        
        # Apply additional replacements
        for para in doc.paragraphs:
            original_text = para.text
            modified_text = original_text
            
            for old_text, new_text in additional_replacements.items():
                if old_text in modified_text:
                    modified_text = modified_text.replace(old_text, new_text)
                    changes_made += 1
                    print(f"   ✅ Added: '{old_text}' → '{new_text}'")
            
            if modified_text != original_text:
                para.text = modified_text
        
        # Save with additional placeholders
        doc.save(output_path)
        
        print(f"✅ Additional placeholders added: {changes_made}")
        return True
        
    except Exception as e:
        print(f"❌ Error adding additional placeholders: {e}")
        return False

def create_enhanced_templates():
    """Create enhanced templates with all necessary placeholders"""
    
    print("🔧 CREATING ENHANCED TEMPLATES")
    print("=" * 40)
    
    templates_to_enhance = [
        {
            'input': 'temp_templates/20250000_E-Mail_Datenblatt_Kundenname.docx',
            'output': 'temp_templates/ENHANCED_extern_template.docx',
            'type': 'extern',
            'description': 'External Customer Template'
        },
        {
            'input': 'temp_templates/20250000_Datenblatt_Kundenname.docx', 
            'output': 'temp_templates/ENHANCED_intern_template.docx',
            'type': 'intern',
            'description': 'Internal Customer Template'
        }
    ]
    
    for template_info in templates_to_enhance:
        input_path = template_info['input']
        output_path = template_info['output']
        template_type = template_info['type']
        description = template_info['description']
        
        print(f"\n📋 Enhancing: {description}")
        
        if not os.path.exists(input_path):
            print(f"❌ Template not found: {input_path}")
            continue
        
        # Step 1: Basic placeholder replacement
        success = enhance_template_with_placeholders(input_path, output_path)
        
        if success:
            # Step 2: Add additional placeholders
            temp_output = output_path.replace('.docx', '_temp.docx')
            add_additional_placeholders(output_path, temp_output, template_type)
            
            # Replace original with enhanced version
            if os.path.exists(temp_output):
                os.replace(temp_output, output_path)
            
            print(f"🎉 Enhanced template ready: {os.path.basename(output_path)}")
        else:
            print(f"❌ Failed to enhance: {description}")
    
    print(f"\n✅ TEMPLATE ENHANCEMENT COMPLETE")
    print(f"📁 Enhanced templates in 'temp_templates/' folder")
    print(f"💡 These templates now have {{placeholder}} syntax for real vault data")

def validate_enhanced_templates():
    """Validate that enhanced templates have proper placeholders"""
    
    print(f"\n🔍 VALIDATING ENHANCED TEMPLATES")
    print("=" * 40)
    
    enhanced_templates = [
        'temp_templates/ENHANCED_extern_template.docx',
        'temp_templates/ENHANCED_intern_template.docx'
    ]
    
    expected_placeholders = [
        '{{primary_email}}',
        '{{primary_email_password}}',
        '{{webmail_url}}',
        '{{customer_name}}',
        '{{smtp_server}}'
    ]
    
    for template_path in enhanced_templates:
        if not os.path.exists(template_path):
            print(f"❌ Enhanced template not found: {os.path.basename(template_path)}")
            continue
        
        print(f"\n📋 Validating: {os.path.basename(template_path)}")
        
        try:
            doc = Document(template_path)
            template_text = ""
            
            # Collect all text
            for para in doc.paragraphs:
                template_text += para.text + " "
            
            # Check for placeholders
            placeholders_found = []
            for placeholder in expected_placeholders:
                if placeholder in template_text:
                    placeholders_found.append(placeholder)
                    print(f"   ✅ Found: {placeholder}")
                else:
                    print(f"   ⚠️ Missing: {placeholder}")
            
            # Check for remaining static text that should be replaced
            static_issues = []
            if 'xxx' in template_text.lower():
                static_issues.append("'xxx' still present")
            if 'kunze-marketing.de' in template_text.lower():
                static_issues.append("static kunze-marketing.de URLs")
            
            print(f"📊 Summary: {len(placeholders_found)}/{len(expected_placeholders)} placeholders")
            
            if static_issues:
                print(f"⚠️ Issues: {', '.join(static_issues)}")
            else:
                print(f"✅ Template properly enhanced")
                
        except Exception as e:
            print(f"❌ Error validating template: {e}")

def main():
    """Main function"""
    
    print("🎯 TEMPLATE PLACEHOLDER ENHANCEMENT")
    print("=" * 45)
    
    # Create enhanced templates
    create_enhanced_templates()
    
    # Validate the results
    validate_enhanced_templates()
    
    print(f"\n🎯 NEXT STEPS:")
    print(f"   1. ✅ Enhanced templates created with {{placeholder}} syntax")
    print(f"   2. 🔧 Test with fixed data extraction logic")
    print(f"   3. 📄 Generate PDFs with real vault data")

if __name__ == "__main__":
    main()